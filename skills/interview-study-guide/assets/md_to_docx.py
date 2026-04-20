"""
Convert a markdown interview study guide to a formatted Word (.docx) document.
Usage: python md_to_docx.py <input.md> [output.docx]
If output path is omitted, writes to same directory as input with .docx extension.
"""

import sys
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colours (all as 6-char hex strings) ─────────────────────────────────────
C_IMEC_BLUE  = '00478A'
C_ACCENT     = '1F5C99'
C_DARK_GREY  = '333333'
C_MID_GREY   = '555555'
C_TH_BG      = '00478A'
C_ALT_ROW    = 'F0F5FA'
C_WHITE      = 'FFFFFF'
C_CODE_BG    = 'F0F0F0'
C_CODE_TEXT  = '222222'
C_HR         = '1F5C99'
C_BORDER     = 'CCCCCC'


def rgb(h):
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ── Low-level XML helpers ────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_para_border(para, color=C_BORDER):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), '4')
        bdr.set(qn('w:space'), '4')
        bdr.set(qn('w:color'), color)
        pBdr.append(bdr)
    pPr.append(pBdr)


def set_para_shading(para, hex_color=C_CODE_BG):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    para._p.get_or_add_pPr().append(shd)


def add_horizontal_rule(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), C_HR)
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)


# ── Inline markup parser ─────────────────────────────────────────────────────

def add_inline_runs(para, text, size=11, color=C_DARK_GREY, bold=False, italic=False):
    """Render **bold**, *italic*, `code`, and [link](url) → display text."""
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # strip link URLs

    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            _plain_run(para, text[pos:m.start()], size, color, bold, italic)
        if m.group(2):
            r = para.add_run(m.group(2)); r.bold = True; r.italic = True
            r.font.size = Pt(size); r.font.color.rgb = rgb(color)
        elif m.group(3):
            r = para.add_run(m.group(3)); r.bold = True
            r.font.size = Pt(size); r.font.color.rgb = rgb(color)
        elif m.group(4):
            r = para.add_run(m.group(4)); r.italic = True
            r.font.size = Pt(size); r.font.color.rgb = rgb(color)
        elif m.group(5):
            r = para.add_run(m.group(5))
            r.font.name = 'Courier New'; r.font.size = Pt(9.5)
            r.font.color.rgb = rgb(C_CODE_TEXT)
        pos = m.end()
    if pos < len(text):
        _plain_run(para, text[pos:], size, color, bold, italic)


def _plain_run(para, text, size, color, bold, italic):
    r = para.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = rgb(color)
    if bold:
        r.bold = True
    if italic:
        r.italic = True


# ── Heading factory ──────────────────────────────────────────────────────────

def add_heading(doc, text, level):
    para = doc.add_paragraph()
    para.paragraph_format.keep_with_next = True
    if level == 1:
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(4)
        r = para.add_run(text); r.bold = True
        r.font.size = Pt(20); r.font.color.rgb = rgb(C_IMEC_BLUE)
    elif level == 2:
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after = Pt(4)
        r = para.add_run(text); r.bold = True
        r.font.size = Pt(15); r.font.color.rgb = rgb(C_IMEC_BLUE)
    elif level == 3:
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(3)
        r = para.add_run(text); r.bold = True
        r.font.size = Pt(13); r.font.color.rgb = rgb(C_ACCENT)
    elif level == 4:
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(2)
        r = para.add_run(text); r.bold = True
        r.font.size = Pt(11.5); r.font.color.rgb = rgb(C_MID_GREY)
    else:
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(2)
        add_inline_runs(para, text, size=11, bold=True)


# ── Table builder ────────────────────────────────────────────────────────────

def add_md_table(doc, rows):
    if len(rows) < 2:
        return
    header = [c.strip() for c in rows[0]]
    data = [[c.strip() for c in r] for r in rows[2:]]  # skip separator row
    cols = len(header)

    table = doc.add_table(rows=1 + len(data), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, text in enumerate(header):
        cell = table.cell(0, i)
        set_cell_bg(cell, C_TH_BG)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(9.5)
        r.font.color.rgb = rgb(C_WHITE)

    # Data rows
    for ri, row in enumerate(data):
        for ci, text in enumerate(row):
            if ci >= cols:
                break
            cell = table.cell(ri + 1, ci)
            if ri % 2 == 1:
                set_cell_bg(cell, C_ALT_ROW)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_inline_runs(p, text, size=9.5)

    doc.add_paragraph()


# ── Code block ───────────────────────────────────────────────────────────────

def add_code_block(doc, lines):
    for line in lines:
        para = doc.add_paragraph()
        set_para_shading(para, C_CODE_BG)
        set_para_border(para)
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
        r = para.add_run(line if line.strip() else ' ')
        r.font.name = 'Courier New'
        r.font.size = Pt(9)
        r.font.color.rgb = rgb(C_CODE_TEXT)


# ── Document setup ───────────────────────────────────────────────────────────

def setup_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = rgb(C_DARK_GREY)
    return doc


# ── Main converter ───────────────────────────────────────────────────────────

def convert(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().splitlines()

    doc = setup_document()

    table_rows = []
    in_table = False
    code_lines = []
    in_code = False
    list_buffer = []  # list of (indent_level, text)

    def flush_list():
        for indent, txt in list_buffer:
            para = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            add_inline_runs(para, txt, size=10.5)
        list_buffer.clear()

    def flush_table():
        nonlocal in_table
        if table_rows:
            add_md_table(doc, table_rows)
        table_rows.clear()
        in_table = False

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith('```'):
            if in_code:
                flush_list()
                add_code_block(doc, code_lines)
                code_lines.clear()
                in_code = False
            else:
                flush_list()
                flush_table()
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^\s*---+\s*$', line):
            flush_list()
            flush_table()
            add_horizontal_rule(doc)
            i += 1
            continue

        # Table row
        if line.startswith('|'):
            flush_list()
            in_table = True
            cells = line.split('|')
            if cells and cells[0].strip() == '':
                cells = cells[1:]
            if cells and cells[-1].strip() == '':
                cells = cells[:-1]
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            flush_table()

        # Heading
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            flush_list()
            add_heading(doc, m.group(2).strip(), len(m.group(1)))
            i += 1
            continue

        # Bullet / list
        m = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if m:
            list_buffer.append((len(m.group(1)) // 2, m.group(2)))
            i += 1
            continue
        else:
            flush_list()

        # Blank line
        if line.strip() == '':
            i += 1
            continue

        # Plain paragraph
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)
        add_inline_runs(para, line.strip())
        i += 1

    flush_list()
    flush_table()
    if in_code and code_lines:
        add_code_block(doc, code_lines)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python md_to_docx.py <input.md> [output.docx]")
        sys.exit(1)

    md_path = sys.argv[1]
    docx_path = sys.argv[2] if len(sys.argv) >= 3 else os.path.splitext(md_path)[0] + '.docx'
    convert(md_path, docx_path)
