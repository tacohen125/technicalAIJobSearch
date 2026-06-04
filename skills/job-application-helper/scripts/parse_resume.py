#!/usr/bin/env python3
"""
parse_resume.py — Three-layer DOCX resume section detector.

Detection approach:
  Layer 1 — Word paragraph styles (Heading 1-9, Title, etc.)
             Catches any resume built with proper Word heading styles.

  Layer 2 — Multi-signal heuristic scoring (6 signals per paragraph):
               bold       +2   all content runs are bold
               allcaps    +2   text is ALL CAPS, or w:caps/w:allCaps flag set
               no_tab     +1   no tab character (rejects Company[TAB]Location lines)
               no_bullet  +1   not a list/bullet paragraph
               short_text +1   text < 50 chars (section labels are concise)
               font_bump  +1   font size ≥ 10% above document baseline
               keyword    +3   text matches a known section keyword

             Decision thresholds:
               keyword path    — keyword fired AND total ≥ 4
               no-keyword path — total ≥ 7  (strong formatting, unusual keyword)

  Layer 3 — Table-cell section headers (modern template detection):
             Many current Word resume templates use single-cell tables as
             visual section separators. The parser walks body elements in
             document order (paragraphs AND tables interleaved) and treats
             a table as a section header when its first row's unique cell
             text matches a known section keyword.

  Approach derived from OpenResume's validated heuristics (open-resume.com),
  adapted for DOCX (they are PDF-based). Both projects' core insight:
  section titles are the only content on their line AND carry double emphasis.

Usage:
  python scripts/parse_resume.py resume.docx              # human-readable
  python scripts/parse_resume.py resume.docx --json       # JSON output
  python scripts/parse_resume.py resume.docx --verbose    # show signal breakdown
"""

import argparse
import json
import re
import sys
import zipfile
from collections import Counter
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.table import Table as _DocxTable
except ImportError:
    print("ERROR: python-docx is required.  pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# Section keyword map
# ---------------------------------------------------------------------------

SECTION_KEYWORDS: dict[str, list[str]] = {
    "SUMMARY":        ["summary", "profile", "objective", "about", "overview",
                       "professional summary", "career summary", "professional profile"],
    "EXPERTISE":      ["areas of expertise", "core competencies", "competencies",
                       "areas of focus", "key competencies"],
    "SKILLS":         ["skills", "technical skills", "tools", "technologies",
                       "capabilities", "technical proficiencies", "technical expertise"],
    "ACCOMPLISHMENTS":["key accomplishments", "accomplishments", "achievements",
                       "highlights", "selected achievements", "career highlights",
                       "key technical achievements", "technical achievements",
                       "key achievements"],
    "EXPERIENCE":     ["experience", "work experience", "employment history",
                       "professional experience", "career history", "work history",
                       "employment",
                       # Leadership / activities sections common on early-career resumes
                       "leadership experience", "leadership activities",
                       "leadership & activities", "activities and leadership",
                       "extracurricular activities", "activities",
                       "volunteer experience", "community involvement"],
    "EDUCATION":      ["education", "academic background", "academics",
                       "educational background", "degrees", "academic credentials"],
    "PUBLICATIONS":   ["publications", "research", "peer-reviewed publications",
                       "papers", "articles", "journal articles",
                       "select publications", "selected publications",
                       "publications and patents", "select publications and patents",
                       "refereed publications", "conference papers"],
    "PRESENTATIONS":  ["presentations", "talks", "conference presentations",
                       "invited talks", "speaking",
                       "select presentations", "selected presentations",
                       "invited presentations", "poster presentations"],
    "PROJECTS":       ["projects", "side projects", "personal projects",
                       "selected projects", "key projects"],
    "CERTIFICATIONS": ["certifications", "licenses", "credentials",
                       "professional certifications", "certificates"],
    "AWARDS":         ["awards", "honors", "recognitions",
                       "awards & honors", "awards and honors", "honors & awards"],
    "VOLUNTEER":      ["volunteer", "volunteering", "community service",
                       "civic engagement", "community involvement"],
    "LANGUAGES":      ["languages", "language skills"],
    "INTERESTS":      ["interests", "hobbies", "activities"],
}

# Flat lookup: normalized keyword → section type
_KEYWORD_LOOKUP: dict[str, str] = {}
for _sec_type, _keywords in SECTION_KEYWORDS.items():
    for _kw in _keywords:
        _KEYWORD_LOOKUP[_kw.lower().strip()] = _sec_type

# Word paragraph style names that identify headings
_HEADING_STYLE_NAMES = {
    "heading 1", "heading 2", "heading 3", "heading 4",
    "heading 5", "heading 6", "title", "subtitle",
}


# ---------------------------------------------------------------------------
# Signal extraction helpers
# ---------------------------------------------------------------------------

def _text(para) -> str:
    """All visible text in a paragraph, stripped."""
    return para.text.strip()


def _any_bold(para) -> bool:
    """
    True if the paragraph contains any explicitly-bold run.
    Checks both run-level rPr and paragraph-level rPr (pPr/w:rPr).
    Uses XML inspection to handle the tri-state (True/False/None=inherit)
    behaviour of python-docx's run.font.bold.
    """
    xml = para._p.xml
    # Find every <w:b> or <w:b/> occurrence
    for m in re.finditer(r'<w:b(?:\s[^>]*)?\s*/?>|<w:b(?:\s[^>]*)?>(?:</w:b>)?', xml):
        tag = m.group(0)
        # Ignore explicit bold-off (w:val="0" or w:val="false")
        if 'w:val="0"' not in tag and 'w:val="false"' not in tag:
            return True
    return False


def _all_runs_bold(para) -> bool:
    """
    True if every non-whitespace content run is explicitly bold.
    Used as a stronger signal than _any_bold when we need higher confidence.
    """
    content_runs = [r for r in para.runs if r.text.strip()]
    if not content_runs:
        return False
    for run in content_runs:
        xml = run._r.xml
        has_b = bool(re.search(r'<w:b(?:\s[^>]*)?\s*/?>', xml))
        is_off = bool(re.search(r'<w:b[^>]*w:val="(?:0|false)"', xml))
        if not has_b or is_off:
            return False
    return True


def _text_is_allcaps(text: str) -> bool:
    """True if the text (letters only) is all uppercase."""
    letters = [c for c in text if c.isalpha()]
    return bool(letters) and all(c.isupper() for c in letters)


def _any_caps_flag(para) -> bool:
    """
    True if the paragraph text is ALL CAPS, OR any run carries w:caps/w:allCaps.
    Handles both explicit formatting and text convention.
    """
    if _text_is_allcaps(_text(para)):
        return True
    xml = para._p.xml
    return bool(re.search(r'<w:(?:caps|allCaps)(?:\s[^/]*)?\s*/?>', xml))


def _has_tab(para) -> bool:
    """
    True if the paragraph contains a tab element or tab character.
    Key anti-false-positive signal: company/title lines use tabs for
    right-aligned dates/locations; section headers never do.
    """
    return '<w:tab' in para._p.xml or '\t' in para.text


def _has_numbering(para) -> bool:
    """True if the paragraph is a bullet or numbered list item."""
    return '<w:numPr>' in para._p.xml


def _get_font_size_pt(para) -> Optional[float]:
    """
    Return the font size in points from the first content run with an
    explicit w:sz value. Returns None if no explicit size is found.
    Half-point units: sz val="20" → 10pt.
    """
    for run in para.runs:
        if not run.text.strip():
            continue
        m = re.search(r'<w:sz w:val="(\d+)"', run._r.xml)
        if m:
            return int(m.group(1)) / 2.0
    return None


def _detect_baseline_font_size(doc) -> float:
    """
    Estimate body font size (in points) as the most common explicit w:sz
    value across all runs in the document. Falls back to 10pt.
    """
    sizes = []
    for para in doc.paragraphs:
        for run in para.runs:
            if not run.text.strip():
                continue
            m = re.search(r'<w:sz w:val="(\d+)"', run._r.xml)
            if m:
                sizes.append(int(m.group(1)) / 2.0)
    if sizes:
        return Counter(sizes).most_common(1)[0][0]
    return 10.0


def _detect_page_count(path: str) -> Optional[int]:
    """
    Try to read the page count from docProps/app.xml inside the .docx ZIP.
    Word writes this after a save; returns None if missing or unreadable.
    """
    try:
        with zipfile.ZipFile(path, 'r') as z:
            if 'docProps/app.xml' in z.namelist():
                xml = z.read('docProps/app.xml').decode('utf-8')
                m = re.search(r'<Pages>(\d+)</Pages>', xml)
                if m:
                    return int(m.group(1))
    except Exception:
        pass
    return None


# ---------------------------------------------------------------------------
# Document-order block traversal (paragraphs + tables interleaved)
# ---------------------------------------------------------------------------

def iter_block_items(doc):
    """
    Yield every top-level paragraph and table in document body order.

    python-docx's doc.paragraphs skips table content entirely.  This generator
    walks the raw XML children so paragraphs and tables are interleaved exactly
    as they appear in the document, enabling correct section detection for
    table-based resume templates.
    """
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.text.paragraph import Paragraph as _Para
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield _Para(child, doc)
        elif isinstance(child, CT_Tbl):
            yield _DocxTable(child, doc)


def _table_section_header_text(table) -> Optional[str]:
    """
    Return section-keyword text if this table acts as a section-header cell, else None.

    Handles three common modern-template patterns:
      • 1×1 pure-header table     — entire cell is a keyword ("EDUCATION")
      • 1×1 header+content table  — first line of cell is a keyword, rest is content
                                    ("Selected Highlights\\nCollaborated with PI...")
      • N-row table, row-0 keyword — first row's unique cells all share the same
                                    keyword ("Core Competencies" repeated across 4 cols)

    Tables where any cell contains contact information (email, phone, URL) are
    treated as name/contact headers and never classified as section headers.
    """
    try:
        rows = table.rows
        if not rows:
            return None

        # Reject name/contact header tables — any cell with contact-pattern text
        # means this is a header bar (name + email), not a section separator.
        for row in rows:
            for cell in row.cells:
                if _CONTACT_PATTERN.search(cell.text):
                    return None

        # Deduplicate horizontally-merged cells using underlying _tc identity
        seen_tc: set[int] = set()
        first_line_texts: list[str] = []
        for cell in rows[0].cells:
            tc_id = id(cell._tc)
            if tc_id in seen_tc:
                continue
            seen_tc.add(tc_id)
            raw = cell.text.strip()
            if raw:
                first_line = next((ln.strip() for ln in raw.split('\n') if ln.strip()), "")
                if first_line:
                    first_line_texts.append(first_line)

        if not first_line_texts:
            return None

        # All unique first-line texts must resolve to the same section type (or be all-caps)
        sec_types: set[str] = set()
        for t in first_line_texts:
            st = _classify_text(t)
            if st:
                sec_types.add(st)
            elif not _text_is_allcaps(t):
                return None  # mixed content → not a clean header

        if len(sec_types) > 1:
            return None  # conflicting keywords in same row

        if sec_types:
            return first_line_texts[0]

        # All-caps but no keyword match — treat as header if short
        if len(first_line_texts) == 1 and len(first_line_texts[0]) < 60:
            return first_line_texts[0]

        return None
    except Exception:
        return None


def _extract_table_content_lines(table) -> list[str]:
    """
    Extract all non-header text lines from a table as a flat list.

    Rows whose unique cells all start with a section keyword are treated as
    header rows and skipped.  For a 1×1 table where the single cell mixes a
    keyword header line with content lines, only the non-keyword lines are
    returned (the keyword line was already recorded as the section header).
    """
    lines: list[str] = []
    seen_tc: set[int] = set()

    for row in table.rows:
        row_texts: list[str] = []
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in seen_tc:
                continue
            seen_tc.add(tc_id)
            t = cell.text.strip()
            if t:
                row_texts.append(t)

        if not row_texts:
            continue

        # Determine whether this is a keyword-only row
        row_is_keyword = all(
            _classify_text(t.split('\n')[0].strip()) is not None
            for t in row_texts
        )
        if row_is_keyword:
            if len(row_texts) == 1:
                # 1×1 table: emit non-keyword lines embedded in the same cell
                for ln in row_texts[0].split('\n'):
                    ln = ln.strip()
                    if ln and _classify_text(ln) is None:
                        lines.append(ln)
            # Multi-cell keyword row: skip entirely
            continue

        for t in row_texts:
            for ln in t.split('\n'):
                ln = ln.strip()
                if ln:
                    lines.append(ln)

    return lines


# ---------------------------------------------------------------------------
# Section keyword classification
# ---------------------------------------------------------------------------

def _classify_text(text: str) -> Optional[str]:
    """
    Return the SECTION_KEYWORDS type for a header text, or None.

    Matching rules (in order):
    1. Exact match after normalisation (lowercase, strip, strip trailing colon)
    2. Strip common decorative prefixes ("select", "selected", "key",
       "additional") and retry exact match — catches "Select Publications",
       "Key Achievements", etc. without enumerating every variant.
    3. Starts-with match within 30% length tolerance — catches minor extensions
       like "Technical Skills & Tools" matching "technical skills".
    """
    normalized = text.lower().strip().rstrip(':').strip()

    # Rule 1: exact match
    if normalized in _KEYWORD_LOOKUP:
        return _KEYWORD_LOOKUP[normalized]

    # Rule 2: strip decorative prefix and retry
    for prefix in ("select ", "selected ", "key ", "additional ", "relevant ", "other "):
        if normalized.startswith(prefix):
            stripped = normalized[len(prefix):]
            if stripped in _KEYWORD_LOOKUP:
                return _KEYWORD_LOOKUP[stripped]

    # Rule 3: starts-with with length tolerance
    for kw, sec_type in _KEYWORD_LOOKUP.items():
        if normalized.startswith(kw) and len(normalized) <= len(kw) * 1.3:
            return sec_type

    return None


# ---------------------------------------------------------------------------
# Layer 1: Word paragraph style detection
# ---------------------------------------------------------------------------

def _get_style_chain(para) -> list[str]:
    """Return style names in inheritance order (most specific first)."""
    chain = []
    style = para.style
    seen = set()
    while style and len(chain) < 10:
        name = style.name.lower()
        if name in seen:
            break
        seen.add(name)
        chain.append(name)
        style = style.base_style
    return chain


def _layer1_is_heading(para) -> bool:
    """
    True if the paragraph's style (or any ancestor) is a recognised heading style.
    Traverses the inheritance chain so custom styles derived from Heading 1 are caught.
    """
    for name in _get_style_chain(para):
        if name in _HEADING_STYLE_NAMES:
            return True
        if re.match(r'heading\s*\d', name):
            return True
    return False


# ---------------------------------------------------------------------------
# Layer 2: Multi-signal heuristic scoring
# ---------------------------------------------------------------------------

_SIGNAL_POINTS = {
    "bold":       2,
    "allcaps":    2,
    "no_tab":     1,
    "no_bullet":  1,
    "short_text": 1,
    "font_bump":  1,
    "keyword":    3,
}


def _compute_signals(para, baseline_pt: float) -> dict[str, bool]:
    """
    Compute all seven heuristic signals for a paragraph.
    Returns dict of signal_name → bool (fired or not).
    """
    text = _text(para)
    sz = _get_font_size_pt(para)
    sec_type = _classify_text(text)

    return {
        "bold":       _all_runs_bold(para),
        "allcaps":    _any_caps_flag(para),
        "no_tab":     not _has_tab(para),
        "no_bullet":  not _has_numbering(para),
        "short_text": 0 < len(text) < 50,
        "font_bump":  (sz is not None) and (sz > baseline_pt * 1.10),
        "keyword":    sec_type is not None,
    }


def _signal_score(signals: dict[str, bool]) -> int:
    return sum(_SIGNAL_POINTS[k] for k, fired in signals.items() if fired)


def _layer2_is_heading(signals: dict[str, bool]) -> bool:
    """
    Two-path decision (see module docstring for rationale):
      keyword path    — keyword fired  AND  total ≥ 4
      no-keyword path — total ≥ 7 (requires strong multi-signal formatting)
    """
    total = _signal_score(signals)
    if signals["keyword"]:
        return total >= 4
    return total >= 7


# ---------------------------------------------------------------------------
# Name and contact info detection
# ---------------------------------------------------------------------------

_CONTACT_PATTERN = re.compile(
    r'[@|]|linkedin|github|http|www\.|'
    r'\d{3}[\s.\-]\d{3}[\s.\-]\d{4}|'  # phone
    r'\b\d{5}\b',                        # ZIP
    re.IGNORECASE,
)

_EMAIL_RE    = re.compile(r'[\w.+\-]+@[\w.\-]+\.\w+')
_PHONE_RE    = re.compile(r'\(?\d{3}\)?[\s.\-]\d{3}[\s.\-]\d{4}')
_LINKEDIN_RE = re.compile(r'linkedin\.com/in/[\w\-]+', re.IGNORECASE)
_LOCATION_RE = re.compile(r'\b[A-Z][a-z]+(?: [A-Z][a-z]+)*,\s+[A-Z]{2}\b')


def detect_contact_info(doc, first_header_idx: int) -> dict:
    """
    Extract contact fields from the pre-header area of the resume.

    Scans the same window as detect_name() and looks for:
      email    — user@domain.tld
      phone    — (NNN) NNN-NNNN or variants
      linkedin — linkedin.com/in/handle
      location — City, ST pattern (not always present)

    Returns a dict with whichever fields were found (keys absent if not found).
    """
    limit = min(first_header_idx, 7) if first_header_idx > 0 else 7
    info: dict = {}
    for para in doc.paragraphs[:limit]:
        text = _text(para)
        if not text:
            continue
        if not info.get("email"):
            m = _EMAIL_RE.search(text)
            if m:
                info["email"] = m.group(0)
        if not info.get("phone"):
            m = _PHONE_RE.search(text)
            if m:
                info["phone"] = m.group(0)
        if not info.get("linkedin"):
            m = _LINKEDIN_RE.search(text)
            if m:
                val = m.group(0)
                info["linkedin"] = val if val.startswith("http") else "https://" + val
        if not info.get("location"):
            m = _LOCATION_RE.search(text)
            if m:
                info["location"] = m.group(0)
    return info


def _detect_name_from_table_cells(pre_header_blocks: list) -> Optional[str]:
    """
    Scan table cells in the pre-header area for a plausible candidate name.

    Modern resume templates often place the candidate's name in a header table
    (e.g. a 1×2 table: "LILLIAN COHEN" | "email@domain.com") rather than in a
    body paragraph.  This is called as a fallback when detect_name() returns
    "Unknown".
    """
    for block in pre_header_blocks:
        if not isinstance(block, _DocxTable):
            continue
        for row in block.rows:
            seen_tc: set[int] = set()
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen_tc:
                    continue
                seen_tc.add(tc_id)
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    if _CONTACT_PATTERN.search(text):
                        continue
                    if text.count('|') >= 2:
                        continue
                    if _classify_text(text) is not None:
                        continue
                    if _is_plausible_name(text):
                        return text
    return None


def _supplement_contact_from_tables(pre_header_blocks: list, info: dict) -> None:
    """
    Fill any missing fields in `info` by scanning table cells in the pre-header area.
    Mutates `info` in place.
    """
    for block in pre_header_blocks:
        if not isinstance(block, _DocxTable):
            continue
        for row in block.rows:
            seen_tc: set[int] = set()
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen_tc:
                    continue
                seen_tc.add(tc_id)
                text = cell.text.strip()
                if not text:
                    continue
                if not info.get("email"):
                    m = _EMAIL_RE.search(text)
                    if m:
                        info["email"] = m.group(0)
                if not info.get("phone"):
                    m = _PHONE_RE.search(text)
                    if m:
                        info["phone"] = m.group(0)
                if not info.get("linkedin"):
                    m = _LINKEDIN_RE.search(text)
                    if m:
                        val = m.group(0)
                        info["linkedin"] = val if val.startswith("http") else "https://" + val
                if not info.get("location"):
                    m = _LOCATION_RE.search(text)
                    if m:
                        info["location"] = m.group(0)


def _is_plausible_name(text: str) -> bool:
    """
    Return True if 'text' looks like a person's name rather than body content.

    Rejects text that is clearly a sentence:
      - More than 6 words
      - More than 60 characters
      - Contains a verb-like lowercase word after the first word
        (heuristic: a real name is Title Case with ≤ 6 tokens)
      - Ends with a full stop (excluding single initials like "J.")
    """
    if len(text) > 60:
        return False
    words = text.split()
    if len(words) > 6:
        return False
    # Reject if any word after the first is all-lowercase and long (sentence word)
    for w in words[1:]:
        clean = w.rstrip('.,;:').lower()
        if len(clean) > 3 and clean == w.rstrip('.,;:'):
            # Word is lowercase and not punctuation-only — looks like prose
            return False
    return True


def _detect_name_from_doc_header(doc) -> Optional[str]:
    """
    Fallback: scan Word document section headers for a plausible name.

    Some older resumes (pre-2020 Word templates) put the candidate's name in
    the page header rather than in the body. This searches `doc.sections[0].header`
    paragraphs for a short, plausible name string.
    """
    try:
        header_paras = doc.sections[0].header.paragraphs
    except Exception:
        return None
    for para in header_paras:
        text = para.text.strip()
        if not text:
            continue
        if _CONTACT_PATTERN.search(text):
            continue
        if text.count('|') >= 2:
            continue
        if _classify_text(text) is not None:
            continue
        if _is_plausible_name(text):
            return text
    return None


def detect_name(doc, first_header_idx: int) -> str:
    """
    Detect the candidate name from the top of the resume (before first header).

    Priority:
    1. Paragraph with the largest explicit font size in the first 7 body lines
       (only if it passes the plausibility check)
    2. First centred body paragraph that passes the plausibility check
    3. First non-empty, non-contact-info body paragraph that passes the check
    4. First plausible paragraph from the Word document page header (older
       resume formats put the name in the header, not the body)
    5. "Unknown" — triggers a manual name prompt in onboard.py
    """
    limit = min(first_header_idx, 7) if first_header_idx > 0 else 7
    candidates = []

    for i, para in enumerate(doc.paragraphs[:limit]):
        text = _text(para)
        if not text or _CONTACT_PATTERN.search(text):
            continue
        # Skip pipe-delimited contact bars  (e.g. "Seattle, WA | 555-1234 | …")
        if text.count('|') >= 2:
            continue
        # Skip paragraphs whose text is a known section keyword (e.g. "Objective",
        # "Summary") — on older resumes the first paragraph may be the section header
        # rather than the candidate's name.
        if _classify_text(text) is not None:
            continue

        sz = _get_font_size_pt(para)
        is_centered = '<w:jc w:val="center"' in para._p.xml

        candidates.append({"text": text, "sz": sz or 0.0, "centered": is_centered})

    if candidates:
        max_sz = max(c["sz"] for c in candidates)
        if max_sz > 0:
            large = [c for c in candidates if c["sz"] == max_sz]
            if _is_plausible_name(large[0]["text"]):
                return large[0]["text"]

        centered = [c for c in candidates if c["centered"]]
        for c in centered:
            if _is_plausible_name(c["text"]):
                return c["text"]

        for c in candidates:
            if _is_plausible_name(c["text"]):
                return c["text"]

    # Fallback: check Word page header (older resume templates)
    from_header = _detect_name_from_doc_header(doc)
    if from_header:
        return from_header

    return "Unknown"


# ---------------------------------------------------------------------------
# Section content extraction
# ---------------------------------------------------------------------------

def _extract_bullets(paras) -> list[str]:
    """Return text of all bulleted/numbered paragraphs."""
    return [_text(p) for p in paras if _has_numbering(p) and _text(p)]


def _extract_skills_text(paras) -> str:
    """Concatenate non-empty paragraph text from a skills section."""
    parts = [_text(p) for p in paras if _text(p)]
    return " | ".join(parts) if parts else ""


def _detect_jobs(paras) -> list[dict]:
    """
    Detect job entries within an Experience section.

    Primary strategy (used when any tab+non-bullet lines are found):
      Resumes commonly use TWO consecutive tab-delimited lines per role:
        Line 1 — Company[TAB]Location      (bold company name)
        Line 2 — Job Title[TAB]Dates       (title, possibly underlined)
      The parser pairs these: the FIRST tab line opens a new job entry;
      the SECOND consecutive tab line (while awaiting a title) sets the title.
      Bold non-bullet lines with no tab are treated as secondary-role entries
      within the same company block (e.g. a role change mid-contract).

    Fallback (no tab lines): bold non-bullet paragraphs under ~80 chars are
    used as job separators (handles simpler single-line-per-role templates).
    """
    jobs: list[dict] = []
    current: Optional[dict] = None
    bullet_count = 0
    awaiting_title = False   # True after a company line, before the title line

    has_tab_pattern = any(
        _has_tab(p) and not _has_numbering(p) and _text(p)
        for p in paras
    )

    def _flush():
        nonlocal current, bullet_count, awaiting_title
        if current is not None:
            current["bullet_count"] = bullet_count
            jobs.append(current)
        current = None
        bullet_count = 0
        awaiting_title = False

    for para in paras:
        text = _text(para)
        if not text:
            continue

        is_bullet = _has_numbering(para)
        has_tab   = _has_tab(para)

        if has_tab_pattern:
            if has_tab and not is_bullet:
                if awaiting_title and current is not None:
                    # Second tab line = Title[TAB]Dates for the current company.
                    # Do NOT flush — this belongs to the entry we just opened.
                    parts = text.split('\t', 1)
                    current["title"] = parts[0].strip()
                    current["dates"] = parts[1].strip() if len(parts) > 1 else ""
                    awaiting_title = False
                else:
                    # First tab line = new Company[TAB]Location entry.
                    _flush()
                    parts = text.split('\t', 1)
                    current = {
                        "company":      parts[0].strip(),
                        "location":     parts[1].strip() if len(parts) > 1 else "",
                        "title":        "",
                        "dates":        "",
                        "bullet_count": 0,
                        "bullet_texts": [],
                    }
                    awaiting_title = True

            elif not has_tab and not is_bullet and _any_bold(para) and len(text) < 100:
                # Bold non-tab non-bullet line: secondary role at the same company
                # (e.g. "Lithography Process Engineer: Sep 2022 – Feb 2024").
                # Flush the current role and open a new entry using this line as the name.
                _flush()
                current = {
                    "company":      text,
                    "location":     "",
                    "title":        "",
                    "dates":        "",
                    "bullet_count": 0,
                    "bullet_texts": [],
                }
                awaiting_title = False

            elif is_bullet:
                bullet_count += 1
                if current is not None:
                    current["bullet_texts"].append(text)
                awaiting_title = False   # bullets started; title line window closed

        else:
            # Fallback: bold, non-bullet, short line = job separator
            if _any_bold(para) and not is_bullet and len(text) < 80:
                _flush()
                current = {
                    "company":      text,
                    "location":     "",
                    "title":        "",
                    "dates":        "",
                    "bullet_count": 0,
                    "bullet_texts": [],
                }
            elif current and is_bullet:
                bullet_count += 1
                current["bullet_texts"].append(text)

    _flush()
    return jobs


# ---------------------------------------------------------------------------
# Main parse function
# ---------------------------------------------------------------------------

def parse_resume(path: str) -> dict:
    """
    Parse a .docx resume and return structured section data.

    Returns a dict with keys:
      name, source_file, baseline_font_pt, sections, warnings

    Each section dict contains:
      header, type, detection, style_name, paragraph_count, bullet_count
      _signal_score, _signals   (prefixed _ = diagnostic, stripped from --json)
      + type-specific keys: jobs (EXPERIENCE), skills_text (SKILLS/EXPERTISE),
        accomplishments (ACCOMPLISHMENTS), count (PUBLICATIONS/PRESENTATIONS)
    """
    doc = Document(path)
    baseline_pt = _detect_baseline_font_size(doc)
    page_count  = _detect_page_count(path)
    warnings: list[str] = []

    # Walk body elements in document order — yields Paragraph and _DocxTable objects
    block_items = list(iter_block_items(doc))

    # ------------------------------------------------------------------
    # Pass 1 — identify section headers from paragraphs (L1/L2) and
    #           table cells (L3).
    # Each record: (block_idx, item, layer_label, signals, header_text)
    # ------------------------------------------------------------------
    header_records: list[tuple[int, object, str, dict, str]] = []

    for i, item in enumerate(block_items):
        if isinstance(item, _DocxTable):
            hdr_text = _table_section_header_text(item)
            if hdr_text:
                signals: dict = {
                    "bold":       False,
                    "allcaps":    _text_is_allcaps(hdr_text),
                    "no_tab":     True,
                    "no_bullet":  True,
                    "short_text": len(hdr_text) < 50,
                    "font_bump":  False,
                    "keyword":    _classify_text(hdr_text) is not None,
                }
                header_records.append((i, item, "layer3_table_cell", signals, hdr_text))
        else:
            text = _text(item)
            if not text:
                continue
            signals = _compute_signals(item, baseline_pt)
            if _layer1_is_heading(item):
                header_records.append((i, item, "layer1_style", signals, text))
            elif _layer2_is_heading(signals):
                header_records.append((i, item, "layer2_heuristic", signals, text))

    # ------------------------------------------------------------------
    # Pass 2 — detect name and contact info from the pre-header area
    # ------------------------------------------------------------------
    first_hdr_block_idx = header_records[0][0] if header_records else len(block_items)
    pre_header_blocks   = block_items[:first_hdr_block_idx]

    # Map paragraph _p element id → index in doc.paragraphs so the existing
    # detect_name / detect_contact_info functions (which take a paragraph index)
    # still work correctly for paragraph-headed resumes.
    para_id_to_idx = {id(p._p): i for i, p in enumerate(doc.paragraphs)}
    first_para_hdr_idx = len(doc.paragraphs)
    for _, item, _, _, _ in header_records:
        if not isinstance(item, _DocxTable) and hasattr(item, '_p'):
            idx = para_id_to_idx.get(id(item._p))
            if idx is not None:
                first_para_hdr_idx = idx
                break

    # Table-cell name detection takes priority: modern templates put the name
    # in a header table rather than a body paragraph, and the paragraph-based
    # heuristic can misfire on education/job lines in those layouts.
    name = (_detect_name_from_table_cells(pre_header_blocks)
            or detect_name(doc, first_para_hdr_idx))
    contact_info = detect_contact_info(doc, first_para_hdr_idx)
    _supplement_contact_from_tables(pre_header_blocks, contact_info)

    # ------------------------------------------------------------------
    # Pass 3 — slice content per section and extract structured data
    # ------------------------------------------------------------------
    sections = []

    for j, (block_idx, hdr_item, layer, signals, header_text) in enumerate(header_records):
        sec_type   = _classify_text(header_text) or "UNKNOWN"
        style_name = "Table" if isinstance(hdr_item, _DocxTable) else hdr_item.style.name

        next_block_idx = header_records[j + 1][0] if j + 1 < len(header_records) else len(block_items)
        content_blocks = block_items[block_idx + 1 : next_block_idx]

        # Content inline with the header table itself (header+content tables, e.g.
        # a skills grid where row 0 is "Core Competencies" and row 1 is the skills).
        header_inline_lines: list[str] = []
        if isinstance(hdr_item, _DocxTable):
            header_inline_lines = _extract_table_content_lines(hdr_item)

        # Split content blocks into paragraphs and tables
        content_paras  = [b for b in content_blocks
                          if not isinstance(b, _DocxTable) and _text(b)]
        content_tables = [b for b in content_blocks if isinstance(b, _DocxTable)]

        # Flatten all content into a single ordered list of text lines
        para_lines  = [_text(p) for p in content_paras]
        table_lines: list[str] = []
        for t in content_tables:
            table_lines.extend(_extract_table_content_lines(t))
        all_lines = [l for l in header_inline_lines + para_lines + table_lines if l]

        bullet_count = sum(1 for p in content_paras if _has_numbering(p))

        entry: dict = {
            "header":          header_text,
            "type":            sec_type,
            "detection":       layer,
            "style_name":      style_name,
            "paragraph_count": len(all_lines),
            "bullet_count":    bullet_count,
            "_signal_score":   _signal_score(signals),
            "_signals":        {k: v for k, v in signals.items()},
        }

        # Type-specific extraction
        if sec_type == "SUMMARY":
            entry["summary_text"] = " ".join(all_lines)
        elif sec_type == "EXPERIENCE":
            # Job entries are almost always in body paragraphs; pass only paras
            # to _detect_jobs so tab-based company/title parsing still works.
            entry["jobs"] = _detect_jobs(content_paras)
        elif sec_type in ("SKILLS", "EXPERTISE"):
            entry["skills_text"]  = " | ".join(all_lines)
            entry["skills_lines"] = all_lines
        elif sec_type == "ACCOMPLISHMENTS":
            # Capture all content regardless of style (see original note).
            entry["accomplishments"] = all_lines
            entry["bullet_count"]    = len(all_lines)
        elif sec_type == "EDUCATION":
            entry["content_lines"] = all_lines
        elif sec_type in ("PUBLICATIONS", "PRESENTATIONS"):
            entry["count"]         = len(all_lines)
            entry["content_lines"] = all_lines

        sections.append(entry)

    if not sections:
        warnings.append(
            "No section headers detected. The resume may use unusual formatting. "
            "Run with --verbose to inspect paragraph-level signal scores."
        )

    return {
        "name":                name,
        "contact_info":        contact_info,
        "source_file":         str(Path(path).name),
        "baseline_font_pt":    baseline_pt,
        "detected_page_count": page_count,
        "sections":            sections,
        "warnings":            warnings,
    }


# ---------------------------------------------------------------------------
# Human-readable output
# ---------------------------------------------------------------------------

def _out(s: str = "") -> None:
    sys.stdout.buffer.write((s + "\n").encode("utf-8", errors="replace"))


def _print_human(result: dict, verbose: bool = False) -> None:
    pc = result.get("detected_page_count")
    pc_str = f"{pc} page{'s' if pc != 1 else ''}" if pc else "unknown (save in Word to embed)"
    _out()
    _out("=== Resume Parse Results " + "=" * 40)
    _out(f"  File            : {result['source_file']}")
    _out(f"  Name            : {result['name']}")
    _out(f"  Page count      : {pc_str}")
    _out(f"  Body font size  : {result['baseline_font_pt']}pt")
    _out(f"  Sections found  : {len(result['sections'])}")
    _out()

    if not result["sections"]:
        _out("  ⚠  No sections detected.")
    else:
        for sec in result["sections"]:
            if sec["detection"] == "layer1_style":
                layer_tag = "[L1-style]"
            elif sec["detection"] == "layer3_table_cell":
                layer_tag = "[L3-table]"
            else:
                layer_tag = f"[L2-score:{sec['_signal_score']}]"
            _out("  " + "─" * 62)
            _out(f"  Header   : \"{sec['header']}\"")
            _out(f"  Type     : {sec['type']}  {layer_tag}  (style: {sec['style_name']})")
            _out(f"  Content  : {sec['paragraph_count']} paragraphs, "
                 f"{sec['bullet_count']} bullets")

            if verbose:
                sigs     = sec["_signals"]
                fired    = [f"{k}(+{_SIGNAL_POINTS[k]})" for k, v in sigs.items() if v]
                notfired = [k for k, v in sigs.items() if not v]
                _out(f"  Signals  : ✓ {', '.join(fired) if fired else 'none'}")
                if notfired:
                    _out(f"           : ✗ {', '.join(notfired)}")

            if "jobs" in sec:
                _out(f"  Jobs ({len(sec['jobs'])}):")
                for job in sec["jobs"]:
                    loc   = f" — {job['location']}" if job["location"] else ""
                    title = f"\n              title: {job['title']}" if job["title"] else ""
                    _out(f"    • {job['company']}{loc}  [{job['bullet_count']} bullets]{title}")

            if "skills_text" in sec and sec["skills_text"]:
                preview = sec["skills_text"][:120]
                suffix  = "..." if len(sec["skills_text"]) > 120 else ""
                _out(f"  Skills   : {preview}{suffix}")

            if "accomplishments" in sec:
                _out(f"  Accomplishments ({len(sec['accomplishments'])}):")
                for acc in sec["accomplishments"]:
                    _out(f"    • {acc[:100]}{'...' if len(acc) > 100 else ''}")

            if sec["type"] in ("PUBLICATIONS", "PRESENTATIONS") and "count" in sec:
                _out(f"  Entries  : {sec['count']}")

        _out("  " + "─" * 62)

    _out()
    if result["warnings"]:
        for w in result["warnings"]:
            _out(f"  ⚠  {w}")
        _out()


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    ap = argparse.ArgumentParser(
        description="Parse a .docx resume and detect section structure.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    ap.add_argument("resume",        help="Path to .docx resume file")
    ap.add_argument("--json",        action="store_true", help="Output JSON")
    ap.add_argument("--verbose", "-v", action="store_true",
                    help="Show per-section signal breakdown")
    args = ap.parse_args()

    path = args.resume
    if not Path(path).exists():
        print(f"ERROR: File not found: {path}", file=sys.stderr)
        sys.exit(1)
    if not path.lower().endswith(".docx"):
        print(f"ERROR: Expected a .docx file, got: {path}", file=sys.stderr)
        sys.exit(1)

    result = parse_resume(path)

    if args.json:
        # Strip internal diagnostic keys from JSON output
        clean_sections = []
        for sec in result["sections"]:
            clean_sections.append({k: v for k, v in sec.items() if not k.startswith("_")})
        print(json.dumps({**result, "sections": clean_sections}, indent=2))
    else:
        _print_human(result, verbose=args.verbose)


if __name__ == "__main__":
    main()
